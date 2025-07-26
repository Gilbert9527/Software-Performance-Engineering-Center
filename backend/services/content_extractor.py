import os
import logging
from typing import Optional, Dict, Any
from dataclasses import dataclass

# File processing libraries
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
    from openpyxl import load_workbook
except ImportError:
    openpyxl = None
    load_workbook = None

try:
    import markdown
except ImportError:
    markdown = None


@dataclass
class ExtractionResult:
    """Result of content extraction"""
    success: bool
    content: str
    error_message: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


class ContentExtractor:
    """Extract text content from various file formats"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are available"""
        missing_deps = []
        
        if PyPDF2 is None:
            missing_deps.append("PyPDF2 (for PDF processing)")
        if Document is None:
            missing_deps.append("python-docx (for Word document processing)")
        if openpyxl is None:
            missing_deps.append("openpyxl (for Excel processing)")
        if markdown is None:
            missing_deps.append("markdown (for Markdown processing)")
        
        if missing_deps:
            self.logger.warning(f"Missing dependencies: {', '.join(missing_deps)}")
    
    def extract_content(self, file_path: str, file_type: str) -> ExtractionResult:
        """
        Extract content from file based on type
        
        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, docx, xlsx, etc.)
            
        Returns:
            ExtractionResult with content and metadata
        """
        if not os.path.exists(file_path):
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"File not found: {file_path}"
            )
        
        try:
            if file_type == 'pdf':
                return self.extract_pdf(file_path)
            elif file_type == 'docx':
                return self.extract_word(file_path)
            elif file_type == 'doc':
                return self.extract_word(file_path)
            elif file_type == 'xlsx':
                return self.extract_excel(file_path)
            elif file_type == 'xls':
                return self.extract_excel(file_path)
            elif file_type == 'md':
                return self.extract_markdown(file_path)
            elif file_type == 'txt':
                return self.extract_text(file_path)
            else:
                return ExtractionResult(
                    success=False,
                    content="",
                    error_message=f"Unsupported file type: {file_type}"
                )
        except Exception as e:
            self.logger.error(f"Error extracting content from {file_path}: {str(e)}")
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"Extraction failed: {str(e)}"
            )
    
    def extract_pdf(self, file_path: str) -> ExtractionResult:
        """
        Extract text content from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            ExtractionResult with extracted text
        """
        if PyPDF2 is None:
            return ExtractionResult(
                success=False,
                content="",
                error_message="PyPDF2 library not available for PDF processing"
            )
        
        try:
            content_parts = []
            metadata = {"pages": 0, "has_text": False}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content_parts.append(f"--- Page {page_num + 1} ---\n{text.strip()}")
                            metadata["has_text"] = True
                    except Exception as e:
                        self.logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                return ExtractionResult(
                    success=False,
                    content="",
                    error_message="No text content found in PDF (may be image-based)",
                    metadata=metadata
                )
            
            return ExtractionResult(
                success=True,
                content=content,
                metadata=metadata
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"PDF extraction failed: {str(e)}"
            )
    
    def extract_word(self, file_path: str) -> ExtractionResult:
        """
        Extract text content from Word document
        
        Args:
            file_path: Path to Word document
            
        Returns:
            ExtractionResult with extracted text
        """
        if Document is None:
            return ExtractionResult(
                success=False,
                content="",
                error_message="python-docx library not available for Word document processing"
            )
        
        try:
            doc = Document(file_path)
            content_parts = []
            metadata = {"paragraphs": 0, "tables": 0}
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
                    metadata["paragraphs"] += 1
            
            # Extract tables
            for table in doc.tables:
                metadata["tables"] += 1
                table_content = []
                
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content.append(cell_text)
                    
                    if row_content:
                        table_content.append(" | ".join(row_content))
                
                if table_content:
                    content_parts.append(f"--- Table ---\n" + "\n".join(table_content))
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                return ExtractionResult(
                    success=False,
                    content="",
                    error_message="No text content found in Word document",
                    metadata=metadata
                )
            
            return ExtractionResult(
                success=True,
                content=content,
                metadata=metadata
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"Word document extraction failed: {str(e)}"
            )
    
    def extract_excel(self, file_path: str) -> ExtractionResult:
        """
        Extract data from Excel file
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            ExtractionResult with extracted data
        """
        if openpyxl is None:
            return ExtractionResult(
                success=False,
                content="",
                error_message="openpyxl library not available for Excel processing"
            )
        
        try:
            workbook = load_workbook(file_path, data_only=True)
            content_parts = []
            metadata = {"sheets": len(workbook.sheetnames), "total_rows": 0, "total_cols": 0}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = []
                
                # Get sheet dimensions
                max_row = sheet.max_row
                max_col = sheet.max_column
                
                if max_row == 1 and max_col == 1 and sheet.cell(1, 1).value is None:
                    continue  # Skip empty sheets
                
                metadata["total_rows"] += max_row
                metadata["total_cols"] = max(metadata["total_cols"], max_col)
                
                sheet_content.append(f"--- Sheet: {sheet_name} ---")
                
                # Extract data row by row
                for row_num in range(1, min(max_row + 1, 1001)):  # Limit to 1000 rows
                    row_data = []
                    has_data = False
                    
                    for col_num in range(1, min(max_col + 1, 51)):  # Limit to 50 columns
                        cell_value = sheet.cell(row_num, col_num).value
                        if cell_value is not None:
                            row_data.append(str(cell_value))
                            has_data = True
                        else:
                            row_data.append("")
                    
                    if has_data:
                        # Remove trailing empty cells
                        while row_data and not row_data[-1]:
                            row_data.pop()
                        
                        if row_data:
                            sheet_content.append(" | ".join(row_data))
                
                if len(sheet_content) > 1:  # More than just the header
                    content_parts.append("\n".join(sheet_content))
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                return ExtractionResult(
                    success=False,
                    content="",
                    error_message="No data found in Excel file",
                    metadata=metadata
                )
            
            return ExtractionResult(
                success=True,
                content=content,
                metadata=metadata
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"Excel extraction failed: {str(e)}"
            )
    
    def extract_markdown(self, file_path: str) -> ExtractionResult:
        """
        Extract content from Markdown file
        
        Args:
            file_path: Path to Markdown file
            
        Returns:
            ExtractionResult with extracted content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                "lines": len(content.splitlines()),
                "characters": len(content),
                "has_markdown": markdown is not None
            }
            
            # If markdown library is available, we could convert to HTML
            # but for AI analysis, raw markdown is often better
            if not content.strip():
                return ExtractionResult(
                    success=False,
                    content="",
                    error_message="Markdown file is empty",
                    metadata=metadata
                )
            
            return ExtractionResult(
                success=True,
                content=content,
                metadata=metadata
            )
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                
                return ExtractionResult(
                    success=True,
                    content=content,
                    metadata={"encoding": "latin-1"}
                )
            except Exception as e:
                return ExtractionResult(
                    success=False,
                    content="",
                    error_message=f"Encoding error: {str(e)}"
                )
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"Markdown extraction failed: {str(e)}"
            )
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """
        Extract content from plain text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            ExtractionResult with extracted content
        """
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                "lines": len(content.splitlines()),
                "characters": len(content),
                "encoding": "utf-8"
            }
            
            if not content.strip():
                return ExtractionResult(
                    success=False,
                    content="",
                    error_message="Text file is empty",
                    metadata=metadata
                )
            
            return ExtractionResult(
                success=True,
                content=content,
                metadata=metadata
            )
            
        except UnicodeDecodeError:
            # Try with different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    
                    return ExtractionResult(
                        success=True,
                        content=content,
                        metadata={"encoding": encoding}
                    )
                except UnicodeDecodeError:
                    continue
            
            return ExtractionResult(
                success=False,
                content="",
                error_message="Could not decode text file with any supported encoding"
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                error_message=f"Text extraction failed: {str(e)}"
            )
    
    def get_supported_formats(self) -> Dict[str, bool]:
        """
        Get supported formats and their availability
        
        Returns:
            Dictionary mapping format to availability
        """
        return {
            'pdf': PyPDF2 is not None,
            'docx': Document is not None,
            'doc': Document is not None,
            'xlsx': openpyxl is not None,
            'xls': openpyxl is not None,
            'md': True,  # Always supported
            'txt': True,  # Always supported
        }
    
    def validate_extraction_capability(self, file_type: str) -> bool:
        """
        Check if extraction is supported for given file type
        
        Args:
            file_type: File type to check
            
        Returns:
            True if extraction is supported
        """
        supported = self.get_supported_formats()
        return supported.get(file_type, False)